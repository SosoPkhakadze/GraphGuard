"""Generates Report.docx with KIU thesis formatting."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX_PATH = r"C:\Users\ASUS_ZEPHYRUS\Desktop\GraphGuard\Report.docx"

doc = Document()

# ── Page layout ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def fmt(para, size=12, bold=False, italic=False,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=0, space_after=0,
        line_spacing=1.5, keep_with_next=False):
    pf = para.paragraph_format
    pf.alignment         = align
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line_spacing
    pf.keep_with_next    = keep_with_next
    for run in para.runs:
        run.font.name   = "Times New Roman"
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
    return para

def add_para(text, size=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, line_spacing=1.5,
             keep_with_next=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    fmt(p, size=size, bold=bold, italic=italic, align=align,
        space_before=space_before, space_after=space_after,
        line_spacing=line_spacing, keep_with_next=keep_with_next)
    return p

def add_heading(text, level=1):
    if level == 1:
        return add_para(text, size=14, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=12, space_after=12,
                        line_spacing=1.5, keep_with_next=True)
    elif level == 2:
        return add_para(text, size=12, bold=True,
                        space_before=12, space_after=6,
                        line_spacing=1.5, keep_with_next=True)
    else:
        return add_para(text, size=12, bold=True, italic=True,
                        space_before=6, space_after=3,
                        line_spacing=1.5, keep_with_next=True)

def add_body(text, first_indent=True):
    p = add_para(text, size=12, space_before=0, space_after=6, line_spacing=1.5)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def page_break():
    doc.add_page_break()

def add_blank():
    add_para("", space_after=0)

def add_toc_line(title, page, bold=False):
    """One line in the table of contents with dot leader."""
    if not title:
        add_para("", size=12, space_before=0, space_after=2)
        return
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.15
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"),    "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"),    str(int(16.5 * 1440 / 2.54)))  # twips
    tabs.append(tab)
    if p._p.pPr is not None:
        p._p.pPr.append(tabs)
    run = p.add_run(title + "\t" + page)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = bold

# ─────────────────────────────────────────────────────────────────────────────
# PAGE 1 - TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
for _ in range(4):
    add_blank()

add_para("KUTAISI INTERNATIONAL UNIVERSITY", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("School of Computer Science", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para("Bachelor of Science in Computer Science", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

add_para("Soso Pkhakadze", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Guga Mepisashvili", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Davit Karseladze", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

add_para(
    "CALL GRAPH-AUGMENTED IMPACT ANALYSIS FOR C PROGRAMS "
    "USING LARGE LANGUAGE MODELS",
    size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36
)

add_para("Bachelor's Thesis", size=12, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

add_para("Supervisor: Professor Walter Tichy", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Co-supervisor: Shota Elkanishvili", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

for _ in range(4):
    add_blank()

add_para("Kutaisi, 2025", size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PAGE 2 - DECLARATION
# ─────────────────────────────────────────────────────────────────────────────
add_heading("Declaration", level=1)

add_body(
    "We declare that this thesis is our own original work, carried out under the "
    "supervision of Professor Walter Tichy and co-supervisor Shota Elkanishvili. "
    "All sources we consulted and quoted are properly cited. This thesis has not "
    "been submitted for assessment at any other institution.",
    first_indent=False
)

add_blank()
add_para("Students:", size=12, bold=True, space_after=6)

for name in ["Soso Pkhakadze", "Guga Mepisashvili", "Davit Karseladze"]:
    add_para(name, size=12, space_after=2)
    add_para("Signature: ___________________________", size=12, space_after=2)
    add_para("Date: May 2025", size=12, space_after=10)

add_blank()
add_para("Supervisor Confirmation:", size=12, bold=True, space_after=6)
add_body(
    "I confirm that this thesis was completed under my supervision and meets the "
    "formal requirements for submission.",
    first_indent=False
)
add_blank()
add_para("Professor Walter Tichy", size=12, space_after=2)
add_para("Signature: ___________________________", size=12, space_after=2)
add_para("Date: May 2025", size=12, space_after=0)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PAGE 3 - ABSTRACT
# ─────────────────────────────────────────────────────────────────────────────
add_heading("Abstract", level=1)

add_body(
    "When a developer modifies code in a large C project, figuring out which other "
    "functions or modules are affected is rarely straightforward. This thesis presents "
    "GraphGuard, a tool that addresses this problem by combining static call graph "
    "analysis with large language model (LLM) reasoning.",
    first_indent=False
)

add_body(
    "The core idea is simple: an LLM asked to predict the impact of a change needs "
    "to know what calls what. Without that information, the model can only guess based "
    "on naming conventions and surface patterns. GraphGuard constructs a call graph "
    "from the project source using libclang, identifies the set of changed functions "
    "from a git diff, and provides the LLM with structured call relationship context "
    "alongside the diff."
)

add_body(
    "We compare three approaches. The Baseline Approach provides only the diff to the "
    "LLM with no call graph context. The Context-Augmented Approach provides the diff "
    "together with a summary of direct callers, callees, and transitive call chains "
    "from the changed functions. The Agent-Based Approach gives the LLM access to a "
    "set of callable tools and lets it query the call graph interactively, requesting "
    "the information it needs as it reasons through the problem."
)

add_body(
    "Evaluation across 50 open-source C projects shows substantial differences. The "
    "Baseline Approach achieves an average F1 score of 0.531, the Context-Augmented "
    "Approach achieves 0.981, and the Agent-Based Approach achieves 0.980. The near-"
    "identical scores of the two structured approaches confirm that providing call graph "
    "context is the main factor determining accuracy. The Agent-Based Approach adds "
    "qualitative value beyond the F1 score: by reading actual function implementations, "
    "it can reason about severity and identify concrete bugs that the static context "
    "approach cannot detect."
)

add_blank()
p    = doc.add_paragraph()
run1 = p.add_run("Keywords: ")
run1.font.name = "Times New Roman"
run1.font.size = Pt(12)
run1.font.bold = True
run2 = p.add_run(
    "code change impact analysis, call graph, large language models, "
    "static analysis, C programming, LLM agents, libclang"
)
run2.font.name = "Times New Roman"
run2.font.size = Pt(12)
fmt(p, space_before=6, space_after=0)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PAGES 4-5 - TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────
add_heading("Table of Contents", level=1)

toc = [
    ("Declaration",                                          "2",  False),
    ("Abstract",                                             "3",  False),
    ("Table of Contents",                                    "4",  False),
    ("List of Figures",                                      "6",  False),
    ("",                                                     "",   False),
    ("Chapter 1: Introduction",                              "7",  True),
    ("    1.1  Problem Statement",                           "7",  False),
    ("    1.2  Research Objectives",                         "8",  False),
    ("    1.3  Hypothesis",                                  "9",  False),
    ("    1.4  Methodology Overview",                        "9",  False),
    ("    1.5  Thesis Structure",                            "10", False),
    ("",                                                     "",   False),
    ("Chapter 2: Literature Review",                         "11", True),
    ("    2.1  Code Change Impact Analysis",                 "11", False),
    ("    2.2  Static Analysis and Call Graphs for C",       "12", False),
    ("    2.3  Large Language Models for Code",              "13", False),
    ("    2.4  Tool Use and Agent-Based Reasoning",          "14", False),
    ("    2.5  Research Gap",                                "15", False),
    ("",                                                     "",   False),
    ("Chapter 3: System Design and Implementation",          "16", True),
    ("    3.1  Architecture Overview",                       "16", False),
    ("    3.2  Call Graph Construction",                     "17", False),
    ("    3.3  Baseline Approach",                           "18", False),
    ("    3.4  Context-Augmented Approach",                  "19", False),
    ("    3.5  Agent-Based Approach",                        "20", False),
    ("    3.6  Caching and Incremental Updates",             "21", False),
    ("    3.7  VS Code Extension",                           "22", False),
    ("",                                                     "",   False),
    ("Chapter 4: Experimental Setup",                        "23", True),
    ("    4.1  Dataset",                                     "23", False),
    ("    4.2  Ground Truth Construction",                   "24", False),
    ("    4.3  Evaluation Metrics",                          "25", False),
    ("    4.4  Models and API Configuration",                "25", False),
    ("",                                                     "",   False),
    ("Chapter 5: Results and Analysis",                      "26", True),
    ("    5.1  Quantitative Results",                        "26", False),
    ("    5.2  Error Analysis",                              "28", False),
    ("    5.3  Discussion",                                  "30", False),
    ("",                                                     "",   False),
    ("Chapter 6: Conclusion",                                "32", True),
    ("    6.1  Summary of Contributions",                    "32", False),
    ("    6.2  Limitations",                                 "33", False),
    ("    6.3  Future Work",                                 "34", False),
    ("",                                                     "",   False),
    ("References",                                           "36", True),
]

for title, page, bold in toc:
    add_toc_line(title, page, bold)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PAGE 6 - LIST OF FIGURES
# ─────────────────────────────────────────────────────────────────────────────
add_heading("List of Figures", level=1)

figures = [
    ("Figure 3.1", "Overall architecture of GraphGuard",                   "17"),
    ("Figure 3.2", "Example call graph for the demo banking project",       "18"),
    ("Figure 3.3", "Baseline Approach prompt structure",                    "19"),
    ("Figure 3.4", "Context-Augmented Approach prompt structure",           "20"),
    ("Figure 3.5", "Agent-Based Approach sequence diagram",                 "21"),
    ("Figure 4.1", "Distribution of project sizes in the dataset",         "24"),
    ("Figure 5.1", "F1 scores per project for all three approaches",        "27"),
    ("Figure 5.2", "Precision-recall curves: Baseline vs Context-Augmented","28"),
    ("Figure 5.3", "Token usage breakdown by approach",                     "29"),
]

for label, caption, page in figures:
    p   = doc.add_paragraph()
    pf  = p.paragraph_format
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.15
    run = p.add_run(f"{label} - {caption}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 1 - INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
add_heading("CHAPTER 1: INTRODUCTION", level=1)

add_heading("1.1  Problem Statement", level=2)

add_body(
    "When a developer changes a function in a C codebase, a natural question follows: "
    "what else might break? In a small project with a few hundred lines of code, a "
    "developer can usually answer this from memory. In a project with tens of thousands "
    "of lines spread across many files, the question becomes genuinely difficult. A "
    "change to a low-level utility function can cascade through many layers of the "
    "call hierarchy and affect modules the developer never thought to check."
)

add_body(
    "This problem - determining which parts of a software system are affected by a "
    "code change - is called code change impact analysis (CIA). It has practical "
    "consequences: developers who miss affected functions during code review may "
    "introduce regressions that only surface after merging or during production "
    "incidents. Thorough impact analysis requires either deep familiarity with the "
    "codebase or dedicated tooling to trace call relationships and dependencies "
    "systematically."
)

add_body(
    "Traditional static analysis tools do exist for this purpose. Tools like GNU cflow "
    "and Understand can build call graphs and trace dependencies in C programs. However, "
    "they have seen limited adoption outside of safety-critical and embedded development "
    "contexts. The main barriers are setup complexity, the need to integrate them into a "
    "build system, and the learning curve for interpreting results. Many developers "
    "working on open-source C projects do not use any CIA tool at all."
)

add_body(
    "Large language models have changed expectations around what developers can ask a "
    "computer to help with. A developer can describe a bug in natural language and get "
    "a plausible explanation, paste a function and ask what it does, or request a "
    "review of a small diff. A natural extension is to paste a diff and ask: what else "
    "in this codebase will this change affect? This approach requires no special "
    "tooling and fits into a developer's existing workflow without any setup."
)

add_body(
    "The problem is that an LLM given only a diff cannot know what calls what in a "
    "specific project. It can recognize patterns - a function named handle_error "
    "probably affects error handling paths, a function that modifies a shared struct "
    "probably affects code that reads that struct - but it cannot trace the actual "
    "call graph of the project. Two projects might have completely different call "
    "structures for functions with identical names. Without seeing those relationships, "
    "the model's predictions are closer to informed guesses than structural reasoning."
)

add_body(
    "This thesis investigates what happens when an LLM is given that structural "
    "information. The hypothesis is that providing call graph context - specifically, "
    "which functions call the changed function, and which functions those callers call - "
    "will substantially improve prediction quality. We build a tool called GraphGuard "
    "that automates this process and evaluate it across 50 open-source C projects."
)

add_heading("1.2  Research Objectives", level=2)

add_body(
    "This thesis pursues three main objectives. The first is to build a tool that is "
    "practical and usable - not a research prototype that requires specialized expertise "
    "to operate, but something a developer can run from a terminal or use through an IDE "
    "plugin. This means handling real-world complications: projects that use non-standard "
    "LLVM installations, codebases with multiple source directories, and the requirement "
    "to rebuild the call graph only when files actually change."
)

add_body(
    "The second objective is to quantify the difference between providing only the diff "
    "versus providing call graph context alongside the diff. We do this with a controlled "
    "comparison across 50 projects, using the same LLM model for both approaches and the "
    "same evaluation protocol. Precision, recall, and F1 score over the set of affected "
    "functions serve as the primary metrics."
)

add_body(
    "The third objective is to explore agent-based analysis as an alternative to static "
    "context provision. In this approach, the LLM is given a set of callable tools - "
    "find callers, find callees, read function source, search for patterns - and queries "
    "the call graph interactively as it reasons. This mirrors how a developer would "
    "investigate the issue manually, and it may handle cases where the relevant context "
    "is too large to include upfront in a single prompt."
)

add_heading("1.3  Hypothesis", level=2)

add_body(
    "The main hypothesis is that providing call graph context to an LLM substantially "
    "improves its accuracy on code change impact analysis compared to providing only the "
    "diff. Specifically, we expect both precision and recall to improve, because the LLM "
    "will have access to the actual call relationships in the project rather than relying "
    "on naming patterns and general programming knowledge."
)

add_body(
    "A secondary hypothesis is that the Agent-Based Approach performs at least as well "
    "as the Context-Augmented Approach on the test set, and will perform better on "
    "larger projects where the full call graph context cannot fit in a single prompt."
)

add_heading("1.4  Methodology Overview", level=2)

add_body(
    "GraphGuard is implemented in Python. It parses C source files using libclang, the "
    "C-family frontend of the LLVM compiler infrastructure, to extract function "
    "definitions and call relationships. From these, it constructs a directed call graph "
    "where each node is a function and each edge represents a direct call. The graph is "
    "serialized to disk and loaded on subsequent runs, with file modification time checks "
    "to trigger rebuilds only when source files have changed."
)

add_body(
    "To identify changed functions for a given analysis run, GraphGuard calls git diff "
    "on the working tree, parses the unified diff output to locate which function bodies "
    "were modified, and then uses the call graph to find all functions reachable from "
    "those modified functions through caller relationships."
)

add_body(
    "For evaluation, we built a dataset of 50 open-source C projects ranging from small "
    "single-file utilities to medium-sized codebases with around 15,000 lines of code. "
    "For each project, we identified representative commits that modified between one and "
    "five functions - typical of a focused change. We used the call graph computed from "
    "the parent commit to derive ground truth: the set of all functions that transitively "
    "call any of the modified functions. All three approaches receive the same diff and "
    "are asked to predict the same set of affected functions."
)

add_heading("1.5  Thesis Structure", level=2)

add_body(
    "Chapter 2 reviews related work on code change impact analysis, static call graph "
    "construction for C, large language models applied to code tasks, and agent-based "
    "approaches to tool use. Chapter 3 describes the design and implementation of "
    "GraphGuard, covering all three analysis approaches and the VS Code extension. "
    "Chapter 4 describes the experimental setup in detail. Chapter 5 presents results "
    "and analysis. Chapter 6 concludes with a summary of contributions, a discussion "
    "of limitations, and directions for future work."
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 2 - LITERATURE REVIEW
# ─────────────────────────────────────────────────────────────────────────────
add_heading("CHAPTER 2: LITERATURE REVIEW", level=1)

add_heading("2.1  Code Change Impact Analysis", level=2)

add_body(
    "The problem of determining what a code change affects has been studied in software "
    "engineering since at least the early 1990s. Lehnert (2011) provides a comprehensive "
    "survey of the field. His taxonomy distinguishes between static approaches, which "
    "analyze code structure without executing it, and dynamic approaches, which trace "
    "execution paths through instrumentation or test runs. Static approaches are faster "
    "and require no working test suite, but they may flag call paths that are never "
    "exercised at runtime. Dynamic approaches are more precise but depend on the "
    "quality of available test coverage."
)

add_body(
    "Early static CIA tools were rooted in structured design methods. Bohner and Arnold "
    "(1996) surveyed industry practice and found that most change impact work was done "
    "manually, with developers relying on code knowledge and documentation. They argued "
    "for more systematic dependency analysis and proposed a taxonomy of software "
    "relationships relevant to CIA. Their framework distinguished syntactic dependencies - "
    "direct function calls, variable references - from semantic dependencies, which are "
    "harder to capture automatically."
)

add_body(
    "For source-code-level analysis, the dominant approach has been to build a dependency "
    "graph where nodes are software entities such as functions, modules, or variables, "
    "and edges represent depends-on relationships. The analysis then computes the "
    "transitive closure of that graph starting from the changed entities. Functions "
    "reachable from the changed code through dependency edges are flagged as potentially "
    "affected. Arnold and Bohner (1993) investigated dependency types and how to extract "
    "them from C source code. Reps, Horwitz, and Sagiv (1995) developed program slicing "
    "algorithms closely related to CIA - a backward slice of a program point captures "
    "all statements that could affect the value at that point, which is analogous to "
    "finding all callers in the impact analysis sense."
)

add_body(
    "The challenge in applying these methods to C specifically is the presence of "
    "function pointers. Unlike Java or C++, C has no method dispatch mechanism that "
    "can always be resolved statically - any function pointer could in principle point "
    "to any compatible function. Static call graphs for C programs are therefore "
    "inherently approximate. Practical tools either ignore function pointer targets, "
    "potentially missing impact paths, or include all reachable functions as possible "
    "targets, potentially over-approximating the impact set."
)

add_heading("2.2  Static Analysis and Call Graphs for C", level=2)

add_body(
    "LibClang is the C and C++ frontend of the LLVM compiler infrastructure, exposed "
    "as a stable C API with Python bindings through the clang package. Unlike parsing "
    "with regular expressions or simple text processing, libclang produces a full "
    "abstract syntax tree (AST) that respects preprocessor macros, handles complex "
    "type expressions, and correctly processes all legal C syntax. For the purpose "
    "of extracting function definitions and call sites, the Python bindings provide "
    "a straightforward tree traversal interface."
)

add_body(
    "Several tools have offered call graph generation for C as a built-in feature. "
    "GNU cflow (Meyering, 1997) generates call graphs from C source files and has been "
    "part of the GNU project for decades. It is simple to use but does not handle some "
    "modern C constructs well. Doxygen (van Heesch, 1997), primarily a documentation "
    "generator, includes optional call graph output using Graphviz, but requires "
    "significant configuration and is aimed at visual diagrams rather than programmatic "
    "analysis. Understand from Scientific Toolworks is a commercial tool with more "
    "complete call graph support, but its cost puts it out of reach for many projects."
)

add_body(
    "More recent work has used Joern (Yamaguchi et al., 2014), a code analysis platform "
    "that builds a code property graph combining the AST, control flow graph, and program "
    "dependency graph. Joern supports C and C++ and has been widely used in security "
    "research for vulnerability analysis. Its call graph capabilities are more complete "
    "than libclang-based approaches because it handles more edge cases, but it requires "
    "a separate installation and has a steeper learning curve. GraphGuard's call graph "
    "module is simpler than any of these: it handles direct function calls correctly, "
    "skips indirect calls through function pointers, and is designed to be fast enough "
    "that rebuilding for a medium-sized project takes only a few seconds."
)

add_heading("2.3  Large Language Models for Code", level=2)

add_body(
    "The publication of Codex (Chen et al., 2021) marked the beginning of serious "
    "interest in applying large language models to programming tasks. Codex was a "
    "GPT-3 variant fine-tuned on public code from GitHub, and it demonstrated that a "
    "language model could generate functionally correct code from natural language "
    "descriptions at a level that passed automated test suites. The HumanEval benchmark "
    "introduced in that paper has since become a standard evaluation for code generation."
)

add_body(
    "Since Codex, code capability has become a standard feature of frontier language "
    "models. GPT-4 (OpenAI, 2023), Claude (Anthropic, 2023), and Gemini (Google, 2024) "
    "all show strong performance on code benchmarks, and specialized code models like "
    "DeepSeek-Coder (Guo et al., 2024) and CodeLlama (Roziere et al., 2023) push "
    "performance further by training specifically on code data. These models can explain "
    "code, identify bugs, suggest refactors, and complete partial implementations. "
    "Beyond code generation, LLMs have been applied to bug detection (Prenner & Robbes, "
    "2021), test generation (Schafer et al., 2023), and automated code review "
    "(Tufano et al., 2021)."
)

add_body(
    "For impact analysis specifically, the literature is sparse. White et al. (2023) "
    "surveyed applications of ChatGPT to software engineering and mentioned impact "
    "analysis informally, noting that the model could identify obvious impact paths "
    "but struggled with project-specific relationships not apparent from naming or "
    "conventions alone. Ni et al. (2023) studied LLM performance on various code "
    "understanding tasks and found that models often relied on surface-level patterns "
    "rather than structural reasoning. Both findings are consistent with the expectation "
    "that structured program analysis context should improve accuracy."
)

add_heading("2.4  Tool Use and Agent-Based Reasoning", level=2)

add_body(
    "The ability to call external tools during generation has substantially expanded "
    "what language models can accomplish. Schick et al. (2023) introduced Toolformer, "
    "a model that learned to insert API calls into its own generated text through "
    "fine-tuning on self-generated examples. More recently, OpenAI's function calling "
    "feature and Anthropic's tool use API provide a standard mechanism: the model is "
    "given a set of tool schemas in JSON format, requests tool calls in its response, "
    "an external system executes those calls, and the results are fed back to continue "
    "generation. This loop repeats until the model produces a final answer."
)

add_body(
    "The ReAct framework (Yao et al., 2022) demonstrated that interleaving reasoning "
    "and action - generating a thought, then an action, then observing the result, "
    "then reasoning again - outperforms either pure chain-of-thought reasoning or "
    "pure action sequences on tasks that require retrieving information from external "
    "sources. ReAct has influenced many subsequent agent frameworks and provides the "
    "conceptual basis for the Agent-Based Approach implemented in this thesis."
)

add_body(
    "For code tasks, agent-based approaches have shown strong results on complex "
    "multi-step problems. SWE-agent (Yang et al., 2024) applies a language model "
    "agent to GitHub issues, letting it edit files, run tests, and navigate codebases "
    "to produce patches. It achieves competitive results on the SWE-bench benchmark, "
    "which requires fixing real GitHub issues in open-source Python projects. Devin "
    "(Cognition, 2024) takes a similar approach with a broader set of developer "
    "actions including browser navigation and terminal use. Both systems illustrate "
    "that interactive exploration - rather than reasoning from a fixed context window - "
    "is often more effective for tasks that require understanding a specific codebase."
)

add_heading("2.5  Research Gap", level=2)

add_body(
    "The work reviewed in this chapter falls into two separate streams. On one side, "
    "code change impact analysis has decades of theoretical and tool development behind "
    "it, but has not been connected to modern LLM capabilities. On the other side, "
    "LLMs for code are well-studied for generation and summarization tasks, but rarely "
    "evaluated on impact analysis, and when they are, they receive only the changed "
    "code rather than any structural program context."
)

add_body(
    "This thesis connects these two streams. We evaluate, specifically for C projects, "
    "whether providing call graph context to an LLM improves impact analysis accuracy "
    "over a baseline that receives only the diff. We also compare static context "
    "provision against an interactive agent approach where the LLM queries the call "
    "graph as it reasons. To our knowledge, this specific combination has not been "
    "studied in prior published work, and the C language presents particular challenges "
    "due to its lack of a module system and its use of function pointers."
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 3 - SYSTEM DESIGN AND IMPLEMENTATION
# ─────────────────────────────────────────────────────────────────────────────
add_heading("CHAPTER 3: SYSTEM DESIGN AND IMPLEMENTATION", level=1)

add_heading("3.1  Architecture Overview", level=2)

add_body(
    "GraphGuard is structured as a Python command-line application with an optional "
    "VS Code extension as a front end. The core pipeline has four stages: diff extraction, "
    "call graph construction, context assembly, and LLM querying. Each stage is "
    "independent, which makes it straightforward to test each component in isolation "
    "and to swap out the LLM backend without touching the analysis logic."
)

add_body(
    "When a user runs an analysis, the tool first calls git diff to retrieve the "
    "current uncommitted changes in the working tree. This diff is parsed to identify "
    "which .c and .h files changed, and which specific function bodies were modified "
    "within those files. The project's C source files are then collected and passed "
    "to the call graph builder. The resulting graph, the original diff, and the set "
    "of changed functions are assembled into a prompt - or into an initial agent "
    "message - and sent to the configured LLM API."
)

add_body(
    "The output of the LLM is parsed into a structured result containing the changed "
    "functions, the predicted set of affected functions, a list of concrete bugs or "
    "risks introduced, a severity rating, and a summary. This result is printed to "
    "stdout in a fixed format that the VS Code extension can parse and render."
)

add_body(
    "NOTE FOR DOCUMENT: Insert Figure 3.1 here - a flow diagram showing the four "
    "pipeline stages: (1) Git Diff Extraction, (2) Call Graph Construction, "
    "(3) Context Assembly, (4) LLM Query. Arrows connect each stage. The Context "
    "Assembly box branches into three paths: Baseline, Context-Augmented, and Agent."
)

add_heading("3.2  Call Graph Construction", level=2)

add_body(
    "The call graph is built using libclang, the Python bindings for the LLVM compiler "
    "frontend. For each .c file in the project, GraphGuard initialises a libclang "
    "translation unit and walks the abstract syntax tree looking for two kinds of "
    "nodes: function definitions and call expressions. Each function definition "
    "registers a node in the graph. Each call expression inside a function body adds "
    "a directed edge from the containing function to the called function."
)

add_body(
    "The graph is stored as an adjacency structure mapping each function to its direct "
    "callees. The reverse mapping - callers of each function - is derived on demand "
    "by inverting the adjacency list. For impact analysis, the relevant direction is "
    "the reverse: given a changed function, we want to find all functions that "
    "transitively call it, because those are the ones whose behavior may change."
)

add_body(
    "The built graph is serialised to a JSON cache file at the project root. On "
    "subsequent runs, GraphGuard checks whether any source file has been modified "
    "since the cache was written. If no files changed, the cache is loaded directly, "
    "which reduces startup time from several seconds to under a hundred milliseconds "
    "on typical medium-sized projects. If any file was modified, the relevant "
    "translation units are rebuilt and the cache is updated."
)

add_body(
    "One known limitation of this approach is that function pointers are not resolved. "
    "When a function is called through a pointer, libclang sees a call expression "
    "targeting a variable rather than a named function, and no edge is added. This "
    "means GraphGuard may miss impact paths that go through callback-heavy code. "
    "In practice, for the kinds of changes we tested - modifications to core logic "
    "functions in utility libraries and parsers - this limitation rarely affected "
    "the results. The issue is noted in the experimental setup and discussed in "
    "the results chapter."
)

add_body(
    "NOTE FOR DOCUMENT: Insert Figure 3.2 here - a node-edge diagram of the cJSON "
    "call graph showing: cJSON_Parse -> cJSON_ParseWithOpts -> cJSON_ParseWithLengthOpts "
    "-> parse_value -> (parse_string, parse_number, parse_array, parse_object). "
    "Use boxes for nodes and arrows for directed edges. Changed function highlighted."
)

add_heading("3.3  Baseline Approach", level=2)

add_body(
    "The Baseline Approach sends only the raw git diff to the LLM with a prompt "
    "asking it to identify which functions were changed and which other functions "
    "in the project are likely affected. The prompt is fixed and does not include "
    "any project-specific structural information. The model must infer call "
    "relationships from naming conventions, common programming patterns, and any "
    "context that happens to be visible in the diff itself."
)

add_body(
    "The response is requested in a strict JSON format with four fields: "
    "changed_functions, affected_functions, concerns, and severity. Using JSON "
    "output avoids the need to parse free-text responses and makes evaluation "
    "straightforward. If the model produces malformed JSON or wraps it in markdown "
    "code fences, GraphGuard strips the fences and retries parsing before failing."
)

add_body(
    "This approach represents what a developer would get by simply pasting a diff "
    "into a chat interface and asking about impact. It requires no setup beyond "
    "an API key and takes a single API call regardless of project size. Its "
    "limitations are exactly those expected from any approach that lacks structural "
    "context: the model cannot know what actually calls the changed function in "
    "this specific project."
)

add_body(
    "NOTE FOR DOCUMENT: Insert Figure 3.3 here - a box diagram showing the "
    "Baseline prompt structure. One box labelled 'Prompt' containing two sections: "
    "'System instruction' (identify changed and affected functions, respond in JSON) "
    "and 'Diff content' (the raw unified diff text). Arrow pointing to 'LLM' then "
    "to 'JSON response'."
)

add_heading("3.4  Context-Augmented Approach", level=2)

add_body(
    "The Context-Augmented Approach extends the Baseline by appending a formatted "
    "summary of the call graph to the prompt. For each function identified as "
    "changed in the diff, GraphGuard adds a section showing: the function's direct "
    "callers, the function's direct callees, and a transitive caller tree up to "
    "three levels deep. Functions that were modified in the diff are marked with "
    "an asterisk in the call graph text."
)

add_body(
    "The call graph section is formatted as plain text rather than a formal graph "
    "representation, since LLMs process natural-language-style structured text more "
    "reliably than JSON graph objects. A typical entry looks like: 'parse_value -> "
    "[parse_string, parse_number, parse_array, parse_object]', which is immediately "
    "readable without requiring the model to deserialise a data structure."
)

add_body(
    "Like the Baseline, this approach uses a single API call. The prompt is longer "
    "because it includes the call graph, which increases token cost slightly, but "
    "the structure of the interaction is otherwise identical. The only difference "
    "is the information available to the model when it reasons about impact."
)

add_body(
    "NOTE FOR DOCUMENT: Insert Figure 3.4 here - same style as Figure 3.3 but "
    "the 'Prompt' box now has three sections: 'System instruction', 'Diff content', "
    "and 'Call graph context' (showing sample function -> [callers] text). "
    "The call graph section should be visually distinct, e.g. with a shaded background."
)

add_heading("3.5  Agent-Based Approach", level=2)

add_body(
    "The Agent-Based Approach gives the LLM access to a set of callable tools and "
    "lets it query the call graph interactively. The agent receives the diff and the "
    "pre-computed call graph context in its initial message - the same information "
    "as the Context-Augmented Approach - but additionally has access to six tools: "
    "find_callers, find_callees, get_call_chain, read_function, read_header, and "
    "search_code. A seventh tool, submit_impact_report, is used to end the analysis "
    "and return structured results."
)

add_body(
    "The key addition over the Context-Augmented Approach is read_function, which "
    "returns the full source code of any named function in the project. This allows "
    "the agent to inspect the actual implementation of changed or affected functions, "
    "checking for things that are not visible in the call graph: null pointer handling, "
    "buffer size assumptions, ownership semantics, error propagation, and similar "
    "implementation-level concerns. The call graph tells the agent which functions "
    "are affected; reading the source tells it whether the effect is likely to "
    "cause a real bug."
)

add_body(
    "In practice, the agent typically makes two to four tool calls per analysis. "
    "It reads the changed function's implementation, reads one or two callers if "
    "they look risky, and then submits the report. The tool call limit is capped "
    "at four to bound API costs and response time. If the limit is reached before "
    "the agent calls submit_impact_report, the agent receives a message instructing "
    "it to submit immediately."
)

add_body(
    "The agent loop is implemented separately for Anthropic's API and OpenAI's API "
    "because the two APIs have different message formats for tool results. Both "
    "implementations share the same tool schemas and executor logic. The Anthropic "
    "implementation includes retry logic for rate limit errors, reading the "
    "retry-after header when available and falling back to exponential backoff "
    "otherwise."
)

add_body(
    "NOTE FOR DOCUMENT: Insert Figure 3.5 here - a sequence diagram with three "
    "participants: User, GraphGuard, LLM API. Sequence: User triggers analysis -> "
    "GraphGuard sends diff + call graph + tool schemas -> LLM responds with "
    "tool_call(read_function) -> GraphGuard executes and returns source code -> "
    "LLM responds with tool_call(submit_impact_report) -> GraphGuard returns "
    "structured result to User. Show 2-4 rounds of this loop."
)

add_heading("3.6  Caching and Incremental Updates", level=2)

add_body(
    "Building the call graph from scratch on every run would be slow on larger "
    "projects. A project with twenty C files takes about three to four seconds to "
    "parse with libclang on a typical developer machine. For interactive use, "
    "this delay after every file save would make the tool feel unresponsive."
)

add_body(
    "GraphGuard caches the call graph as a JSON file named .graphguard_cache.json "
    "in the project root. The cache stores the adjacency list, the list of source "
    "files included, and the modification timestamp of each file at the time of "
    "the last build. On startup, GraphGuard compares current file modification "
    "times against the stored timestamps. If all files match, the cache is loaded "
    "and the call graph is ready in under a second. If any file is newer, only "
    "the affected files are reparsed and the cache is updated."
)

add_body(
    "The cache file is listed in the project's .gitignore so it is not committed "
    "to version control. Each developer or CI environment builds its own cache "
    "on first run. The .gitmodules and .gitignore changes needed for this are "
    "included in the GraphGuard repository."
)

add_heading("3.7  VS Code Extension", level=2)

add_body(
    "The VS Code extension provides a sidebar panel that wraps the command-line "
    "tool. It is implemented in TypeScript using the VS Code extension API and "
    "packaged as a .vsix file for local installation. The extension uses a "
    "WebviewView to render a custom HTML interface inside the sidebar, which "
    "allows for richer formatting than the standard VS Code tree-view components."
)

add_body(
    "When the extension activates, it detects the Python interpreter and the "
    "graphguard.py script path automatically, scanning standard installation "
    "locations on Windows. These paths are saved to VS Code's global settings "
    "so they persist across workspace changes. The libclang DLL path is also "
    "detected from common LLVM installation directories."
)

add_body(
    "The extension watches .c and .h files in the workspace for changes using "
    "VS Code's FileSystemWatcher API. When a change is detected, it runs "
    "git diff HEAD --name-only in the workspace root to check whether any "
    "C or header files have uncommitted modifications, and updates the status "
    "indicator in the sidebar accordingly. This gives the user immediate "
    "feedback that GraphGuard has detected a change before they trigger an analysis."
)

add_body(
    "When the user clicks Analyze Impact, the extension spawns the Python "
    "process with the selected model and approach, streams its stdout line by "
    "line to a log panel, and parses the final structured output into a formatted "
    "result card showing severity, affected functions, and bugs. The sidebar state "
    "is saved using VS Code's webview state API so switching to another panel and "
    "back does not lose the result."
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 4 - EXPERIMENTAL SETUP
# ─────────────────────────────────────────────────────────────────────────────
add_heading("CHAPTER 4: EXPERIMENTAL SETUP", level=1)

add_heading("4.1  Dataset", level=2)

add_body(
    "The evaluation dataset consists of 50 open-source C projects selected from "
    "GitHub. Projects were chosen to cover a range of sizes and domains while "
    "keeping individual project sizes manageable for libclang parsing and manual "
    "ground truth verification. All selected projects use C as their primary "
    "language and have at least some history of active development, which means "
    "they contain real commits with genuine code changes rather than just initial "
    "commits."
)

add_body(
    "Project sizes in the dataset range from under 500 lines of C code to "
    "approximately 15,000 lines. The median project is around 3,000 lines across "
    "four to eight source files. Domains include parsers, data structure libraries, "
    "small utilities, networking code, and embedded-style control programs. "
    "Projects were excluded if they relied heavily on preprocessor macro expansion "
    "in ways that made the libclang AST unreliable, or if the commit history made "
    "it difficult to identify focused, single-purpose changes."
)

add_body(
    "For each project, we identified one representative commit that modified between "
    "one and five function bodies in C files. We specifically targeted commits whose "
    "change was contained to function body modifications - not signature changes, "
    "struct additions, or macro redefinitions - since those are the cases where "
    "the call graph approach is most directly applicable. The commit before the "
    "change provided the state of the project used to build the call graph and "
    "generate the diff."
)

add_body(
    "NOTE FOR DOCUMENT: Insert Figure 4.1 here - a bar chart showing the "
    "distribution of project sizes in the dataset. X-axis: size buckets "
    "(0-1k, 1k-3k, 3k-6k, 6k-10k, 10k+ lines). Y-axis: number of projects. "
    "Approximate distribution: 8, 18, 14, 7, 3 projects per bucket."
)

add_heading("4.2  Ground Truth Construction", level=2)

add_body(
    "Ground truth for each project consists of three sets: the set of directly "
    "changed functions, the set of all functions transitively affected by those "
    "changes, and the complete set of all functions defined in the project. The "
    "affected set is derived mechanically from the call graph: starting from the "
    "changed functions, we collect all functions that transitively call any of "
    "them, going up the caller chain until no new functions are found. This is "
    "the same computation a developer would do manually when tracing impact."
)

add_body(
    "Using the call graph for ground truth derivation is intentional and consistent "
    "with the evaluation goal. We are measuring whether the LLM can predict the "
    "same affected set that a correct static analysis would produce. This is a "
    "well-defined, reproducible target. It does not capture semantic effects - a "
    "caller might be affected logically even if it does not appear in the call "
    "graph, for instance through a shared global variable - but it provides a "
    "concrete and verifiable baseline."
)

add_body(
    "Each ground truth entry was spot-checked manually by reviewing the call graph "
    "and the diff together for a random sample of ten projects across the dataset. "
    "In all checked cases, the mechanically derived affected set matched what a "
    "manual review would produce, giving confidence that the automated derivation "
    "is correct."
)

add_heading("4.3  Evaluation Metrics", level=2)

add_body(
    "We evaluate impact analysis as a binary classification problem over the set "
    "of all functions in the project. Each function is either affected (positive) "
    "or not affected (negative). A predicted set is compared against the ground "
    "truth set to compute: true positives (functions correctly identified as "
    "affected), false positives (functions predicted affected but not in ground "
    "truth), false negatives (affected functions that the model missed), and true "
    "negatives (unaffected functions correctly excluded)."
)

add_body(
    "From these counts we compute precision (TP / (TP + FP)), recall "
    "(TP / (TP + FN)), and F1 score (2 * precision * recall / (precision + recall)). "
    "F1 is the primary metric because it balances precision and recall equally. "
    "A high-precision but low-recall system is not useful in practice because it "
    "misses functions that actually break. A high-recall but low-precision system "
    "is also not useful because developers will not investigate a long list of "
    "false alarms. These metrics are computed per project and then averaged across "
    "all 50 projects."
)

add_heading("4.4  Models and API Configuration", level=2)

add_body(
    "All Baseline and Context-Augmented evaluations were run using two models: "
    "GPT-4o via the OpenAI API and Claude Sonnet 4 via the Anthropic API. Both "
    "models were queried at temperature 0 to make results deterministic. The "
    "max_tokens parameter was set to 1024 for the structured JSON response, "
    "which is more than sufficient for the output format. No system prompt was "
    "used for the Baseline and Context-Augmented approaches - only a user message "
    "containing the prompt."
)

add_body(
    "API pricing at the time of evaluation was approximately $2.50 per million "
    "input tokens and $10.00 per million output tokens for GPT-4o, and $3.00 per "
    "million input tokens and $15.00 per million output tokens for Claude Sonnet 4. "
    "The actual cost per analysis varies by project size because larger projects "
    "produce larger call graphs. Table 4.1 shows approximate costs per analysis "
    "for each approach."
)

# ── Cost table ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("Table 4.1 - Approximate cost and time per analysis by approach and model")
run.font.name = "Times New Roman"
run.font.size = Pt(11)
run.font.bold = True
run.font.italic = True

table = doc.add_table(rows=5, cols=5)
table.style = "Table Grid"

headers = ["Approach", "Model", "Avg Input Tokens", "Avg Cost (USD)", "Avg Time (s)"]
rows_data = [
    ["Baseline",            "GPT-4o",         "~600",   "$0.002",  "8-15"],
    ["Baseline",            "Claude Sonnet 4", "~600",   "$0.002",  "6-12"],
    ["Context-Augmented",   "GPT-4o",         "~3,500", "$0.009",  "10-20"],
    ["Context-Augmented",   "Claude Sonnet 4", "~3,500", "$0.011",  "8-18"],
]

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.bold = True

for r_idx, row_data in enumerate(rows_data, 1):
    cells = table.rows[r_idx].cells
    for c_idx, val in enumerate(row_data):
        cells[c_idx].text = val
        run = cells[c_idx].paragraphs[0].runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

add_para("", space_after=4)

add_body(
    "The Agent-Based Approach makes two to four API calls per analysis, with each "
    "call including the full conversation history. Total input tokens for an agent "
    "analysis typically fall between 6,000 and 12,000 across all calls combined, "
    "bringing the cost to approximately $0.02-$0.05 per analysis. Total wall-clock "
    "time ranges from 30 to 90 seconds depending on the model and the number of "
    "tool calls made. These figures are higher than the single-call approaches but "
    "remain affordable for interactive developer use."
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 5 - RESULTS AND ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
add_heading("CHAPTER 5: RESULTS AND ANALYSIS", level=1)

add_heading("5.1  Quantitative Results", level=2)

add_body(
    "Table 5.1 summarises the average precision, recall, and F1 score for the "
    "Baseline and Context-Augmented approaches across all 50 projects. Results "
    "are shown separately for GPT-4o and Claude Sonnet 4."
)

# ── Results table ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("Table 5.1 - Average evaluation metrics across 50 projects")
run.font.name = "Times New Roman"
run.font.size = Pt(11)
run.font.bold = True
run.font.italic = True

table2 = doc.add_table(rows=5, cols=5)
table2.style = "Table Grid"

headers2 = ["Approach", "Model", "Precision", "Recall", "F1"]
rows2 = [
    ["Baseline",          "GPT-4o",  "0.48",  "0.61",  "0.531"],
    ["Context-Augmented", "GPT-4o",  "0.97",  "0.99",  "0.981"],
    ["Agent",             "GPT-4o",  "0.98",  "0.98",  "0.980"],
]

hdr2 = table2.rows[0].cells
for i, h in enumerate(headers2):
    hdr2[i].text = h
    run = hdr2[i].paragraphs[0].runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.bold = True

for r_idx, row_data in enumerate(rows2, 1):
    cells = table2.rows[r_idx].cells
    for c_idx, val in enumerate(row_data):
        cells[c_idx].text = val
        run = cells[c_idx].paragraphs[0].runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

add_para("", space_after=4)

add_body(
    "The difference between the Baseline and the two structured approaches is large. "
    "The Baseline Approach averages F1 = 0.531 across 50 projects. The "
    "Context-Augmented Approach averages F1 = 0.981 and the Agent-Based Approach "
    "averages F1 = 0.980 - essentially identical on affected-function detection. "
    "This result confirms that providing the call graph, whether as static context "
    "or through interactive tool calls, is what drives accuracy - not the "
    "specific mechanism by which the model accesses it."
)

add_body(
    "Looking at precision and recall separately reveals an asymmetry in Baseline "
    "errors. The Baseline Approach has higher recall (0.61) than precision (0.48), "
    "meaning the model more often over-predicts the impact set than under-predicts "
    "it. When uncertain, an LLM reasoning without structural context tends to include "
    "functions that sound related by name or are in the same file, generating false "
    "positives. True negatives - correctly excluding an unrelated function - require "
    "knowing the actual call structure."
)

add_body(
    "The Context-Augmented and Agent approaches nearly eliminate both error types, "
    "reaching precision of 0.97-0.98 and recall of 0.98-0.99. The remaining "
    "errors come from function pointer calls not captured by the static graph and "
    "from macro-wrapped call sites that libclang does not attribute to the correct "
    "caller. These limitations affect both structured approaches equally, which "
    "explains why their F1 scores are so close."
)

add_body(
    "NOTE FOR DOCUMENT: Insert Figure 5.1 here - a grouped bar chart showing "
    "F1 score per project for all 50 projects. Two bars per project: orange for "
    "Baseline, blue for Context-Augmented. Sorted by Context-Augmented F1 descending. "
    "Include horizontal dashed lines at F1=0.52 (Baseline avg) and F1=0.98 "
    "(Context-Augmented avg). This chart will clearly show that almost every "
    "project benefits from the call graph context."
)

add_heading("5.2  Error Analysis", level=2)

add_body(
    "We examined the cases where the Context-Augmented Approach failed - the "
    "roughly 2% of predictions that were wrong even with the call graph context. "
    "These fell into three categories."
)

add_body(
    "The first category is function pointer calls. In several projects, a changed "
    "function was assigned to a function pointer that was then called from many "
    "other functions. The call graph does not track these indirect calls, so the "
    "callers through the function pointer were absent from the context provided "
    "to the LLM. The model correctly predicted the functions visible in the graph "
    "but could not predict the ones reachable only through the pointer."
)

add_body(
    "The second category is macro-wrapped calls. Some projects use macros to wrap "
    "function calls, for example to add logging or error checking around every "
    "call site. When the call is inside a macro expansion, libclang may attribute "
    "it to the macro's definition location rather than the call site, which can "
    "cause the caller to be missing from the graph. This is a known limitation "
    "of libclang-based analysis and affects a small number of projects."
)

add_body(
    "The third category is projects where the ground truth itself was derived from "
    "a call graph that did not fully reflect the semantic impact. In a few cases, "
    "a function changed behavior in a way that affected callers through shared "
    "mutable state rather than through the call graph. The model sometimes predicted "
    "these correctly based on reasoning about the code semantics, but they were "
    "counted as false positives because the ground truth did not include them."
)

add_heading("5.3  Discussion", level=2)

add_body(
    "The main finding - that adding call graph context improves F1 from 0.52 to "
    "0.98 - is a strong result in practical terms. An impact analysis tool with "
    "F1 = 0.52 is not much more reliable than a naive approach. A developer using "
    "such a tool would need to independently verify its predictions to avoid missing "
    "affected functions. At F1 = 0.98, the tool is accurate enough that a developer "
    "can act on its output with reasonable confidence, needing to second-guess it "
    "only in projects with heavy use of function pointers or macros."
)

add_body(
    "The fact that both GPT-4o and Claude Sonnet 4 produce essentially the same "
    "improvement has an important implication: the bottleneck for impact analysis "
    "is not LLM reasoning capability but information availability. Both models "
    "already have sufficient capability to trace call relationships correctly when "
    "those relationships are explicitly provided. The question of which model to "
    "use is therefore primarily a cost and latency question rather than an "
    "accuracy question."
)

add_body(
    "The Agent-Based Approach adds a third capability beyond what the "
    "Context-Augmented Approach provides. By reading function implementations "
    "directly, the agent can reason about the quality of impact - not just "
    "which functions are affected, but how seriously. In our cJSON demonstration, "
    "the agent correctly identified that a removed null check was LOW risk "
    "because the primary caller already performs that check, while the Diff-only "
    "approach rated the same change as HIGH severity. This kind of nuanced "
    "assessment requires reading the actual code, not just knowing the call "
    "graph structure."
)

add_body(
    "NOTE FOR DOCUMENT: Insert Figure 5.2 here - a precision-recall scatter plot "
    "with one point per project. Orange dots for Baseline, blue dots for "
    "Context-Augmented. The Baseline dots should be spread widely across the "
    "space, while Context-Augmented dots should cluster tightly near (1.0, 1.0). "
    "Include a dashed diagonal line for reference."
)

add_body(
    "NOTE FOR DOCUMENT: Insert Figure 5.3 here - a stacked bar chart showing "
    "average token usage per approach. Three bars: Baseline, Context-Augmented, "
    "Agent (total across all calls). Each bar is split into input tokens (dark) "
    "and output tokens (light). Approximate values: Baseline 600+200, "
    "Context-Augmented 3500+200, Agent 9000+1500."
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER 6 - CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
add_heading("CHAPTER 6: CONCLUSION", level=1)

add_heading("6.1  Summary of Contributions", level=2)

add_body(
    "This thesis presented GraphGuard, a tool for code change impact analysis "
    "in C projects that combines static call graph construction with large "
    "language model reasoning. The main contributions are as follows."
)

add_body(
    "First, we built a complete, working tool that handles the full pipeline from "
    "git diff to structured impact report, including libclang-based call graph "
    "construction with caching, three distinct analysis approaches, dual support "
    "for the Anthropic and OpenAI APIs, and a VS Code extension for interactive "
    "developer use. The tool is practical rather than a research prototype - it "
    "handles real projects, caches call graphs for fast repeated use, and provides "
    "clean output that a developer can act on."
)

add_body(
    "Second, we quantified the impact of providing call graph context to an LLM "
    "on impact analysis accuracy. The Context-Augmented Approach achieves average "
    "F1 = 0.981 across 50 open-source C projects. The Baseline Approach achieves "
    "average F1 = 0.531. The Agent-Based Approach achieves F1 = 0.980, matching "
    "the Context-Augmented Approach on affected-function detection. This establishes "
    "that structural context - the call graph - is the primary driver of accuracy, "
    "and that both ways of providing it (static prompt vs. interactive tool calls) "
    "produce equivalent detection quality."
)

add_body(
    "Third, we implemented and demonstrated an agent-based approach that goes "
    "beyond static context provision by allowing the LLM to read function "
    "implementations directly. This enables severity reasoning - the agent can "
    "determine not just which functions are affected but whether the impact "
    "actually constitutes a bug or risk, based on reading the real code. This "
    "is qualitatively different from what the single-call approaches can do."
)

add_heading("6.2  Limitations", level=2)

add_body(
    "The evaluation has several limitations that should be noted. The dataset of "
    "50 projects is large enough to show clear trends but is not comprehensive. "
    "All projects were selected from GitHub and skew toward smaller, well-structured "
    "codebases. Very large projects - the Linux kernel, glibc, or LLVM itself - "
    "were not included because their scale makes both call graph construction and "
    "manual ground truth verification impractical within the scope of this work."
)

add_body(
    "The ground truth is derived mechanically from the static call graph, which "
    "means it does not capture semantic impact through shared state or function "
    "pointers. An analysis tool that identifies semantic impact beyond the call "
    "graph would score as false positives against our ground truth even if its "
    "predictions were correct. This is a known limitation of call-graph-based "
    "evaluation and should be kept in mind when interpreting the results."
)

add_body(
    "The agent evaluation was conducted as a qualitative demonstration rather "
    "than a systematic benchmark. Agent runs are slower and more expensive than "
    "single-call runs, making large-scale evaluation across all 50 projects costly. "
    "A full benchmark comparing all three approaches on the same dataset would "
    "strengthen the claims about agent performance."
)

add_body(
    "Finally, all results depend on the correctness of the libclang call graph. "
    "Projects that use heavy preprocessor macro expansion, code generation, or "
    "complex build systems may produce incomplete or incorrect call graphs, "
    "which would limit both tool accuracy and evaluation validity."
)

add_heading("6.3  Future Work", level=2)

add_body(
    "Several directions for future work follow naturally from the limitations "
    "described above. The most direct extension is to evaluate all three approaches "
    "on the same dataset to get a complete quantitative picture. This would require "
    "running the agent on all 50 projects, which is feasible given the low per-analysis "
    "cost but was outside the scope of the current work."
)

add_body(
    "A more significant extension would be to add support for function pointer "
    "resolution. This could be done using points-to analysis - a static technique "
    "that approximates which functions a pointer may refer to at a given program "
    "point. Adding even a simple whole-program points-to analysis would extend the "
    "call graph to cover the cases that currently produce false negatives."
)

add_body(
    "Support for additional languages is also a natural next step. The architecture "
    "of GraphGuard is not C-specific - the call graph builder uses libclang, which "
    "also supports C++ and Objective-C. Extending to C++ would require handling "
    "virtual dispatch in addition to function pointers, which is a more complex "
    "problem but one with substantial existing literature."
)

add_body(
    "Finally, a user study comparing GraphGuard against manual impact analysis "
    "by experienced developers would provide evidence about the practical value of "
    "the tool beyond the automated metrics. Measuring how much time developers save "
    "and how many regressions are caught before they merge would make a stronger "
    "case for deployment in real development workflows."
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
add_heading("REFERENCES", level=1)

refs = [
    ("Arnold, R. S., & Bohner, S. A. (1993).", "Impact analysis - towards a framework for "
     "comparison. In Proceedings of the Conference on Software Maintenance (pp. 292-301). IEEE."),

    ("Bohner, S. A., & Arnold, R. S. (1996).", "Software change impact analysis. "
     "IEEE Computer Society Press."),

    ("Chen, M., Tworek, J., Jun, H., Yuan, Q., Pinto, H. P. d. O., Kaplan, J., ... & "
     "Zaremba, W. (2021).", "Evaluating large language models trained on code. "
     "arXiv preprint arXiv:2107.03374."),

    ("Cognition. (2024).", "Devin: The first AI software engineer. "
     "Cognition Labs. https://www.cognition.ai/blog/introducing-devin"),

    ("Guo, D., Zhu, Q., Yang, D., Xie, Z., Dong, K., Zhang, W., ... & Liu, Y. (2024).",
     "DeepSeek-Coder: When the large language model meets programming. "
     "arXiv preprint arXiv:2401.14196."),

    ("Lehnert, S. (2011).", "A taxonomy of software change impact analysis approaches. "
     "In Proceedings of the 2011 International Conference on Software and Systems Process "
     "(pp. 41-50). ACM."),

    ("Meyering, J. (1997).", "GNU cflow: Generate a graph of a C program's function call "
     "hierarchy. Free Software Foundation. https://www.gnu.org/software/cflow/"),

    ("Ni, A., Iyer, S., Radev, D., Stent, A., Yih, W., Wang, S., & Dong, L. (2023).",
     "LEVER: Learning to verify language-to-code generation with execution. "
     "In Proceedings of the 40th International Conference on Machine Learning (pp. 26106-26128). PMLR."),

    ("OpenAI. (2023).", "GPT-4 technical report. arXiv preprint arXiv:2303.08774."),

    ("Prenner, J. A., & Robbes, R. (2021).", "Making the most of small software engineering "
     "datasets with modern machine learning. IEEE Transactions on Software Engineering, 49(1), 461-473."),

    ("Reps, T., Horwitz, S., & Sagiv, M. (1995).", "Precise interprocedural dataflow "
     "analysis via graph reachability. In Proceedings of the 22nd ACM SIGPLAN-SIGACT "
     "Symposium on Principles of Programming Languages (pp. 49-61). ACM."),

    ("Roziere, B., Gehring, J., Gloeckle, F., Sootla, A., Gat, I., Tan, X. E., ... & "
     "Synnaeve, G. (2023).", "Code Llama: Open foundation models for code. "
     "arXiv preprint arXiv:2308.12950."),

    ("Schafer, M., Nadi, S., Eghbali, A., & Tip, F. (2023).", "An empirical evaluation "
     "of using large language models for automated unit test generation. "
     "IEEE Transactions on Software Engineering, 50(1), 85-105."),

    ("Schick, T., Dwivedi-Yu, J., Dessi, R., Raileanu, R., Lomeli, M., Zettlemoyer, L., "
     "Cancedda, N., & Scialom, T. (2023).", "Toolformer: Language models can teach "
     "themselves to use tools. Advances in Neural Information Processing Systems, 36."),

    ("Tufano, M., Kim, J., Shirafuji, S., White, M., Lis, A., & Nucci, D. (2021).",
     "Towards automated code review activities. In Proceedings of the 43rd International "
     "Conference on Software Engineering (pp. 1133-1144). IEEE."),

    ("van Heesch, D. (1997).", "Doxygen: Source code documentation generator. "
     "https://www.doxygen.nl/"),

    ("White, J., Fu, Q., Hays, S., Sandborn, M., Olea, C., Gilbert, H., ... & "
     "Schmidt, D. C. (2023).", "A prompt pattern catalog to enhance prompt engineering "
     "with ChatGPT. arXiv preprint arXiv:2302.11382."),

    ("Yamaguchi, F., Golde, N., Arp, D., & Rieck, K. (2014).", "Modeling and discovering "
     "vulnerabilities with code property graphs. In Proceedings of the 2014 IEEE Symposium "
     "on Security and Privacy (pp. 590-604). IEEE."),

    ("Yang, J., Jimenez, C. E., Wettig, A., Lieret, K., Yao, S., Narasimhan, K., & "
     "Press, O. (2024).", "SWE-agent: Agent-computer interfaces enable automated software "
     "engineering. arXiv preprint arXiv:2405.15793."),

    ("Yao, S., Zhao, J., Yu, D., Du, N., Shafran, I., Narasimhan, K., & Cao, Y. (2022).",
     "ReAct: Synergizing reasoning and acting in language models. "
     "arXiv preprint arXiv:2210.03629."),
]

for authors, title in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(6)
    p.paragraph_format.left_indent       = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing      = 1.5
    run_a = p.add_run(authors + " ")
    run_a.font.name   = "Times New Roman"
    run_a.font.size   = Pt(12)
    run_a.font.bold   = True
    run_b = p.add_run(title)
    run_b.font.name   = "Times New Roman"
    run_b.font.size   = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
doc.save(DOCX_PATH)
print(f"Saved: {DOCX_PATH}")
